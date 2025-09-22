"""
Titanic – Klasifikacija
Autor: Ognjen Ikrašev RA 238/2022
Opis:
    - Učitava isključivo train.csv (Kaggle Titanic format)
    - Radi EDA (grafici i tabele), obradu nedostajućih vrednosti, inženjering atributa
    - Enkodira kategorijske atribute preko OneHotEncoder-a
    - Trenira i poredi više modela (LogisticRegression, RandomForest, GradientBoosting, HistGradientBoosting)
    - Radi k-unakrsnu validaciju sa metrikama: accuracy, precision, recall, f1
    - Vizuelizuje: distribucije, preživljavanje po atributima, značajnost atributa (feature importances & permutation importance)
    - Upoređuje model sa svim atributima vs. samo najbitniji atributi
    - Kreira PNG grafikone i (opciono) .docx izveštaj (ako je instaliran python-docx)

Upotreba:
    1) Postavite train.csv u isti folder.
    2) Pokrenite:  python Titanic_Klasifikacija_train_only.py
    3) Rezultati će se sačuvati u ./outputs/ (grafici, tabele, izveštaj)

"""

import os
import re
import warnings
from dataclasses import dataclass
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder
from sklearn.impute import SimpleImputer
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, classification_report
from sklearn.model_selection import StratifiedKFold, cross_validate, train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.inspection import permutation_importance
from sklearn.feature_selection import SelectFromModel, mutual_info_classif

try:
    from sklearn.ensemble import HistGradientBoostingClassifier

    HGB_AVAILABLE = True
except Exception:
    HGB_AVAILABLE = False

warnings.filterwarnings("ignore", category=FutureWarning)
plt.rcParams.update({"figure.dpi": 120, "figure.autolayout": True})

OUTPUT_DIR = "outputs"
FIG_DIR = os.path.join(OUTPUT_DIR, "figures")
TABLE_DIR = os.path.join(OUTPUT_DIR, "tables")
for d in [OUTPUT_DIR, FIG_DIR, TABLE_DIR]:
    os.makedirs(d, exist_ok=True)


@dataclass
class ModelResult:
    name: str
    metrics: Dict[str, float]
    fitted_pipeline: Pipeline
    feature_names: List[str]


def load_data(path: str = "train.csv") -> pd.DataFrame:
    if not os.path.exists(path):
        raise FileNotFoundError(
            f"Nije pronađen fajl '{path}'. Postavite train.csv u radni direktorijum."
        )
    df = pd.read_csv(path)
    return df


def extract_title(name: str) -> str:
    # Izvlačenje titule iz imena (npr. "Mr.", "Mrs.", "Master.")
    if pd.isna(name):
        return "Unknown"
    match = re.search(r",\s*([^\.]+)\.", name)
    title = match.group(1).strip() if match else "Unknown"
    # Grupisanje retkih titula
    common = {"Mr", "Mrs", "Miss", "Master"}
    return title if title in common else "Rare"


def ticket_prefix(ticket: str) -> str:
    if pd.isna(ticket):
        return "UNKNOWN"
    t = ticket.replace(".", "").replace("/", "").upper()
    parts = [p for p in t.split() if not p.isdigit()]
    return parts[0] if parts else "NUMERIC"


def cabin_deck(cabin: str) -> str:
    if pd.isna(cabin) or not isinstance(cabin, str) or len(cabin) == 0:
        return "U"
    return cabin[0].upper()


def engineer_features(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    # Izvedene promenljive
    df["Title"] = df["Name"].apply(extract_title)
    df["FamilySize"] = df["SibSp"].fillna(0) + df["Parch"].fillna(0) + 1
    df["IsAlone"] = (df["FamilySize"] == 1).astype(int)
    df["TicketPrefix"] = df["Ticket"].apply(ticket_prefix)
    df["CabinDeck"] = df["Cabin"].apply(cabin_deck)

    # Binning za Age (robustnije za nedostajuće vrednosti)
    # Napomena: pravu imputaciju Age radimo u pipeline-u,
    # a ovde definišemo i diskretizovan atribut radi potencijalne informativnosti
    age_bins = [0, 12, 18, 25, 35, 50, 65, 120]
    df["AgeBin"] = pd.cut(df["Age"], bins=age_bins,
                          labels=["child", "teen", "yadult", "adult", "mature", "senior", "elder"], include_lowest=True)

    # Cena karte po log-skali (stabilizuje raspodelu)
    df["FareLog"] = np.log1p(df["Fare"].fillna(0))

    return df


def basic_eda_plots(df: pd.DataFrame, save: bool = True) -> None:
    # 1) Raspodela starosti (hist)
    plt.figure()
    df["Age"].dropna().plot(kind="hist", bins=30)
    plt.title("Raspodela starosti (Age)")
    plt.xlabel("Godine")
    if save:
        plt.savefig(os.path.join(FIG_DIR, "raspodela_starosti.png"))
    plt.close()

    # 2) Preživljavanje po polu
    plt.figure()
    df["Sex_label"] = df["Sex"].map({"female": "zena", "male": "muskaraca"})
    df["Survived_label"] = df["Survived"].map({0: "Nije preživelo", 1: "Preživelo"})
    pd.crosstab(df["Sex_label"], df["Survived_label"]).plot(kind="bar", stacked=True)
    plt.title("Preživljavanje po polu (Sex)")
    plt.xlabel("Pol")
    plt.ylabel("Broj putnika")
    if save:
        plt.savefig(os.path.join(FIG_DIR, "prezivljanje_po_polu.png"))
    plt.close()

    # 3) Preživljavanje po klasi
    plt.figure()
    df["Survived_label"] = df["Survived"].map({0: "Nije preživelo", 1: "Preživelo"})
    pd.crosstab(df["Pclass"], df["Survived_label"]).plot(kind="bar", stacked=True)
    plt.title("Preživljavanje po klasi (Pclass)")
    plt.xlabel("Klasa")
    plt.ylabel("Broj putnika")
    if save:
        plt.savefig(os.path.join(FIG_DIR, "prezivljavanje_po_klasi.png"))
    plt.close()

    # 4) Preživljavanje po luci ukrcavanja
    plt.figure()
    df["Survived_label"] = df["Survived"].map({0: "Nije preživelo", 1: "Preživelo"})
    pd.crosstab(df["Embarked"], df["Survived_label"]).plot(kind="bar", stacked=True)
    plt.title("Preživljavanje po luci (Embarked)")
    plt.xlabel("Luka ukrcavanja")
    plt.ylabel("Broj putnika")
    if save:
        plt.savefig(os.path.join(FIG_DIR, "prezivljavanje_po_luci.png"))
    plt.close()


def build_preprocessor(feature_df: pd.DataFrame) -> Tuple[ColumnTransformer, List[str], List[str]]:
    # Definišemo koje kolone odbacujemo (neinformativne / ID, duplikati informacije itd.)
    drop_cols = ["PassengerId", "Name", "Ticket", "Cabin", "Survived"]

    # Numeričke i kategorijske kolone nakon inženjeringa
    num_cols = [c for c in feature_df.columns if c not in drop_cols]
    # privremeno podela – OneHot će raditi i nad binovima
    num_base = ["Age", "SibSp", "Parch", "Fare", "FamilySize", "IsAlone", "FareLog"]
    cat_cols = ["Pclass", "Sex", "Embarked", "Title", "TicketPrefix", "CabinDeck", "AgeBin"]
    num_cols = [c for c in num_cols if c in num_base]

    # Imputacija i enkodovanje
    numeric_transformer = Pipeline(steps=[
        ("imputer", SimpleImputer(strategy="median")),
    ])

    categorical_transformer = Pipeline(steps=[
        ("imputer", SimpleImputer(strategy="most_frequent")),
        ("onehot", OneHotEncoder(handle_unknown="ignore", sparse_output=False))
    ])

    preprocessor = ColumnTransformer(
        transformers=[
            ("num", numeric_transformer, num_cols),
            ("cat", categorical_transformer, cat_cols),
        ],
        remainder="drop",
    )

    return preprocessor, num_cols, cat_cols


def evaluate_models(X: pd.DataFrame, y: pd.Series, preprocessor: ColumnTransformer) -> List[ModelResult]:
    models = [
        ("LogReg", LogisticRegression(max_iter=1000, n_jobs=None, solver="lbfgs")),
        ("RF", RandomForestClassifier(n_estimators=400, max_depth=None, random_state=42)),
        ("GBoost", GradientBoostingClassifier(random_state=42)),
    ]
    if HGB_AVAILABLE:
        models.append(("HistGB", HistGradientBoostingClassifier(random_state=42)))

    results: List[ModelResult] = []

    scoring = {
        "accuracy": "accuracy",
        "precision": "precision",
        "recall": "recall",
        "f1": "f1"
    }

    skf = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)

    for name, clf in models:
        pipe = Pipeline(steps=[("pre", preprocessor), ("clf", clf)])
        cv = cross_validate(pipe, X, y, cv=skf, scoring=scoring, return_estimator=True)
        metrics = {m: float(np.mean(cv[f"test_{m}"])) for m in scoring.keys()}
        # Fit final pipeline on full data for feature names
        pipe.fit(X, y)
        # Dohvati imena kolona posle transformacije
        ohe = pipe.named_steps["pre"].named_transformers_["cat"].named_steps["onehot"]
        cat_feature_names = list(ohe.get_feature_names_out())
        feature_names = pipe.named_steps["pre"].get_feature_names_out()
        # get_feature_names_out vrati num i cat exp naziv; uprosti imena
        feature_names = [fn.replace("num__", "").replace("cat__", "") for fn in feature_names]

        results.append(ModelResult(name=name, metrics=metrics, fitted_pipeline=pipe, feature_names=feature_names))

        # Sačuvaj tabelu metrika po foldovima
        fold_table = pd.DataFrame({
            "fold": list(range(1, len(cv["fit_time"]) + 1)),
            "accuracy": cv["test_accuracy"],
            "precision": cv["test_precision"],
            "recall": cv["test_recall"],
            "f1": cv["test_f1"],
        })
        fold_table.to_csv(os.path.join(TABLE_DIR, f"cv_folds_{name}.csv"), index=False)

    return results


def plot_feature_importance(model_res: ModelResult, X: pd.DataFrame, y: pd.Series, top_k: int = 20) -> pd.DataFrame:
    pipe = model_res.fitted_pipeline
    clf = pipe.named_steps["clf"]

    # Pokušaj nativnih importanci (drveće i boosting)
    importances = None
    if hasattr(clf, "feature_importances_"):
        importances = clf.feature_importances_
    elif hasattr(clf, "coef_"):
        # Logistic regression – uzmi aps vrednosti koeficijenata
        coefs = clf.coef_.ravel()
        importances = np.abs(coefs) / (np.linalg.norm(coefs) + 1e-9)

    if importances is not None:
        n = min(len(model_res.feature_names), len(importances))
        fi_df = pd.DataFrame({"feature": model_res.feature_names, "importance": importances})
        fi_df = fi_df.sort_values("importance", ascending=False).head(top_k)
        plt.figure()
        plt.barh(fi_df["feature"][::-1], fi_df["importance"][::-1])
        plt.title(f"{model_res.name} – Značajnost atributa (native)")
        plt.xlabel("Importance")
        plt.tight_layout()
        plt.savefig(os.path.join(FIG_DIR, f"feature_importance_native_{model_res.name}.png"))
        plt.close()
    else:
        fi_df = pd.DataFrame()

    # Permutation importance (stabilnije tumačenje)
    # Radimo na train split-u (bez posebnog valida jer radimo CV već iznad)
    X_train, X_val, y_train, y_val = train_test_split(X, y, stratify=y, test_size=0.2, random_state=42)
    pipe.fit(X_train, y_train)
    perm = permutation_importance(pipe, X_val, y_val, n_repeats=15, random_state=42, scoring="f1")

    # Uzimamo stvarne kolone iz pipeline-a
    try:
        feature_names = pipe[:-1].get_feature_names_out()
    except AttributeError:
        feature_names = X.columns

    perm_df = pd.DataFrame({"feature": feature_names, "importance": perm.importances_mean})
    perm_df = perm_df.sort_values("importance", ascending=False).head(top_k)

    plt.figure()
    plt.barh(perm_df["feature"][::-1], perm_df["importance"][::-1])
    plt.title(f"{model_res.name} – Permutation importance (F1)")
    plt.xlabel("ΔF1 pri permutaciji")
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, f"feature_importance_permutation_{model_res.name}.png"))
    plt.close()

    return perm_df


def select_top_features(preprocessor: ColumnTransformer, X: pd.DataFrame, y: pd.Series, k: int = 20) -> List[str]:
    # Dobij sve vektorisane kolone
    preprocessor.fit(X)
    feat_names = preprocessor.get_feature_names_out()
    feat_names = [fn.replace("num__", "").replace("cat__", "") for fn in feat_names]

    # Transformiši X da bi računali MI (zahteva numeričke podatke)
    X_enc = preprocessor.transform(X)
    mi = mutual_info_classif(X_enc, y, random_state=42)
    order = np.argsort(mi)[::-1]

    top_idx = order[: min(k, len(order))]
    top_features = [feat_names[i] for i in top_idx]

    # Sačuvaj tabelu MI
    mi_df = pd.DataFrame({"feature": feat_names, "mutual_info": mi}).sort_values("mutual_info", ascending=False)
    mi_df.to_csv(os.path.join(TABLE_DIR, "mutual_information.csv"), index=False)

    return top_features


def train_with_selected_features(model_name: str, base_pre: ColumnTransformer, X: pd.DataFrame, y: pd.Series,
                                 selected_features: List[str]) -> ModelResult:
    # Rekonstruisanje ColumnTransformer-a da zadrži samo izabrane kolone posle transformacije.
    # Jednostavnije: koristimo postojeći preprocessor i zatim SelectFromModel na RandomForest-u
    # ili, pošto već imamo listu imena nakon transformacije, uradićemo mapiranje maskom.
    # Implementiraćemo custom step koji filtrira kolone po imenu.

    from sklearn.base import BaseEstimator, TransformerMixin

    class ColumnSelectorByName(BaseEstimator, TransformerMixin):
        def __init__(self, names: List[str]):
            self.names = names
            self._all: List[str] = []
            self.idx_: np.ndarray | None = None

        def fit(self, X, y=None):
            if hasattr(X, "columns"):
                self._all = list(X.columns)
            return self

        def transform(self, X):
            # X je numpy nakon preprocesiranja; imena dolaze iz preprocesora
            return X[:, self._indices]

        def set_feature_names(self, all_names: List[str]):
            self._all = all_names
            self._indices = np.array([self._all.index(n) for n in self.names if n in self._all])

    # Napravi kopiju preprocesora da dobijemo imena posle fit-a
    pre = base_pre
    pre.fit(X)
    all_names = [fn.replace("num__", "").replace("cat__", "") for fn in pre.get_feature_names_out()]

    selector = ColumnSelectorByName(selected_features)
    selector.set_feature_names(all_names)

    # Model
    clf = RandomForestClassifier(n_estimators=400, random_state=42)

    pipe = Pipeline([
        ("pre", pre),
        ("sel", selector),
        ("clf", clf)
    ])

    skf = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)
    scoring = {"accuracy": "accuracy", "precision": "precision", "recall": "recall", "f1": "f1"}
    cv = cross_validate(pipe, X, y, cv=skf, scoring=scoring, return_estimator=True)
    metrics = {m: float(np.mean(cv[f"test_{m}"])) for m in scoring.keys()}

    # Fit final
    pipe.fit(X, y)

    # Pošto posle selekcije gubimo originalna imena, dodelićemo samo izabrana
    model_res = ModelResult(name=f"RF_top{len(selected_features)}", metrics=metrics, fitted_pipeline=pipe,
                            feature_names=selected_features)

    # Sačuvaj fold metrike
    fold_table = pd.DataFrame({
        "fold": list(range(1, len(cv["fit_time"]) + 1)),
        "accuracy": cv["test_accuracy"],
        "precision": cv["test_precision"],
        "recall": cv["test_recall"],
        "f1": cv["test_f1"],
    })
    fold_table.to_csv(os.path.join(TABLE_DIR, f"cv_folds_{model_res.name}.csv"), index=False)

    return model_res


def save_summary_docx(results: List[ModelResult], best_model: ModelResult, perm_df: pd.DataFrame) -> None:
    # Pokušaj kreiranja Word izveštaja
    try:
        from docx import Document
        from docx.shared import Inches
        doc = Document()
        doc.add_heading('Titanic – Klasifikacija (train.csv)', 0)
        doc.add_paragraph(
            'Modeli su trenirani isključivo nad train.csv skupom podataka. Slede ključni nalazi i rezultati.')

        doc.add_heading('Metrike po modelima (5-fold CV)', level=1)
        for res in results:
            p = doc.add_paragraph()
            p.add_run(f"{res.name}: ").bold = True
            p.add_run(f"accuracy={res.metrics['accuracy']:.3f}, precision={res.metrics['precision']:.3f}, "
                      f"recall={res.metrics['recall']:.3f}, f1={res.metrics['f1']:.3f}")

        doc.add_heading('Grafici – EDA', level=1)
        for fig in ["age_distribution.png", "survival_by_sex.png", "survival_by_pclass.png",
                    "survival_by_embarked.png"]:
            path = os.path.join(FIG_DIR, fig)
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(5.5))

        doc.add_heading('Značajnost atributa', level=1)
        for fig in os.listdir(FIG_DIR):
            if fig.startswith("feature_importance_"):
                path = os.path.join(FIG_DIR, fig)
                doc.add_picture(path, width=Inches(5.5))

        doc.add_heading('Tumačenje i preporuke', level=1)
        doc.add_paragraph(
            "Najveći doprinos preživljavanju, u skladu sa literaturom i analizom, pokazuju pol (female), klasa karte (Pclass), "
            "kao i starost. Takođe, izvedeni atributi poput FamilySize/IsAlone i titula (Title) doprinose razdvajanju klasa. "
            "Preporuke: dodatno optimizovati hiperparametre (npr. pretraga mrežom), probati kalibraciju verovatnoća i balansiranje klasa."
        )

        out_path = os.path.join(OUTPUT_DIR, "report.docx")
        doc.save(out_path)
    except Exception as e:
        # Fallback: Markdown
        lines = [
            "# Titanic – Klasifikacija (train.csv)",
            "\n## Metrike po modelima (5-fold CV)\n",
        ]
        for res in results:
            lines.append(
                f"- {res.name}: accuracy={res.metrics['accuracy']:.3f}, precision={res.metrics['precision']:.3f}, recall={res.metrics['recall']:.3f}, f1={res.metrics['f1']:.3f}")
        lines.append("\n## Tumačenje i preporuke\n")
        lines.append(
            "Najvažniji faktori: pol (female), klasa (Pclass), starost; korisni izvedeni: FamilySize/IsAlone, Title. ")
        lines.append(
            "Preporuke: dodatna optimizacija hiperparametara, kalibracija verovatnoća, balansiranje klasa, probati i drugačije imputacije starosti.")
        with open(os.path.join(OUTPUT_DIR, "report.md"), "w", encoding="utf-8") as f:
            f.write("\n".join(lines))


def main():
    # 1) Učitavanje i inženjering
    df = load_data("train.csv")
    df = engineer_features(df)

    # 2) EDA grafici
    basic_eda_plots(df)

    # 3) Priprema X, y
    y = df["Survived"].astype(int)
    X = df.drop(columns=["Survived"])

    preprocessor, num_cols, cat_cols = build_preprocessor(df)

    # 4) Evaluacija različitih modela (CV)
    results = evaluate_models(X, y, preprocessor)

    # Izbor najboljeg po F1 (možete promeniti kriterijum)
    best = max(results, key=lambda r: r.metrics["f1"])

    # 5) Značajnost atributa – nativna + permutaciona
    perm_df = plot_feature_importance(best, X, y, top_k=20)
    perm_df.to_csv(os.path.join(TABLE_DIR, f"top_permutation_importance_{best.name}.csv"), index=False)

    # 6) Selekcija najbitnijih atributa po MI i poređenje performansi
    top_feats = select_top_features(preprocessor, X, y, k=20)
    res_top = train_with_selected_features("RF", preprocessor, X, y, top_feats)
    results.append(res_top)

    # 7) Sačuvaj pregled metrika
    metrics_table = pd.DataFrame([
        {"model": r.name, **r.metrics} for r in results
    ]).sort_values("f1", ascending=False)
    metrics_table.to_csv(os.path.join(TABLE_DIR, "model_metrics.csv"), index=False)

    # 8) Izveštaj
    save_summary_docx(results, best, perm_df)

    # 9) Kratak ispis u konzoli
    print("\n===== Metrike po modelima (5-fold CV) =====")
    for r in results:
        print(
            f"{r.name:10s} | acc={r.metrics['accuracy']:.3f}  prec={r.metrics['precision']:.3f}  rec={r.metrics['recall']:.3f}  f1={r.metrics['f1']:.3f}")
    print("\nGrafici i tabele sačuvani su u ./outputs/")


main()
